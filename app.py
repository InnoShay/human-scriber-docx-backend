from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from io import BytesIO

app = Flask(__name__)

def align_map(value: str):
    if not value:
        return WD_ALIGN_PARAGRAPH.LEFT
    v = value.lower()
    if v == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if v == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT

def page_num_format(page, total, style):
    if style == 1:
        return f"{page}"
    if style == 2:
        return f"Page {page}"
    if style == 3:
        return f"{page} of {total}"
    if style == 4:
        return f"Page {page} of {total}"
    return f"{page}"

@app.route("/prepare_docx", methods=["POST"])
def prepare_docx():
    try:
        data = request.get_json()

        text = data.get("text", "")

        margins = data.get("margins", {})
        header = data.get("header", {})
        footer1 = data.get("footer1", {})
        footer2 = data.get("footer2", {})
        cover = data.get("cover_page", {})

        doc = Document()

        section = doc.sections[0]
        sectPr = section._sectPr

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')

        for border in ["top", "bottom", "left", "right"]:
            elem = OxmlElement(f"w:{border}")
            elem.set(qn('w:val'), 'single')      
            elem.set(qn('w:sz'), '8')            
            elem.set(qn('w:space'), '24')        
            elem.set(qn('w:color'), '000000')    
            pgBorders.append(elem)

        sectPr.append(pgBorders)


        section = doc.sections[0]
        section.top_margin = Inches(float(margins.get("top", 1)))
        section.bottom_margin = Inches(float(margins.get("bottom", 1)))
        section.left_margin = Inches(float(margins.get("left", 1)))
        section.right_margin = Inches(float(margins.get("right", 1)))

        if cover:
            cover_title = cover.get("title", "")
            submitted_to = cover.get("submitted_to", "")
            submitted_by = cover.get("submitted_by", "")

            if cover_title or submitted_to or submitted_by:

                for _ in range(10):  
                    doc.add_paragraph()

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cover_title)
                run.bold = True
                run.font.size = Pt(32)

                doc.add_paragraph()

                if submitted_to:
                    p = doc.add_paragraph(f"Submitted To: {submitted_to}")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.size = Pt(18)

                if submitted_by:
                    p = doc.add_paragraph(f"Submitted By: {submitted_by}")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.size = Pt(18)

                doc.add_page_break()

        if header.get("content"):
            for section in doc.sections:
                hdr = section.header.paragraphs[0]
                hdr.text = header["content"]
                hdr.alignment = align_map(header.get("alignment"))

        if text:
            for line in text.split("\n"):
                p = doc.add_paragraph(line)
                p.style = "Normal"

        total_pages = 1  
        fmt = int(footer2.get("format", 1))

        for section in doc.sections:
            ftr_paras = section.footer.paragraphs

            if footer1.get("content"):
                p1 = ftr_paras[0]
                p1.text = footer1["content"]
                p1.alignment = align_map(footer1.get("alignment", "left"))

            if footer2.get("page_num"):
                p2 = section.footer.add_paragraph()
                p2.text = page_num_format(1, total_pages, fmt)

                if footer1.get("alignment") == footer2.get("alignment"):
                    if footer2["alignment"] == "right":
                        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif footer2["alignment"] == "left":
                        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p2.alignment = align_map(footer2.get("alignment"))

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="generated.docx"
        )

    except Exception as e:
        print("Docx generation error:", e)
        return jsonify({"error": "Generation failed", "details": str(e)}), 500
